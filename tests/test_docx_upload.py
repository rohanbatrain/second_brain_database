#!/usr/bin/env python3
"""
Test DOCX upload to the RAG system
"""

import requests
from pathlib import Path
import io

# Load token
token = Path("rag_token.txt").read_text().strip()
headers = {"Authorization": f"Bearer {token}"}

# Create a simple DOCX using python-docx (if available)
try:
    from docx import Document
    from docx.shared import Pt
    
    # Create a new Document
    doc = Document()
    doc.add_heading('Test RAG Upload Document', 0)
    
    p = doc.add_paragraph('This is a test DOCX document for the RAG system.')
    
    doc.add_heading('Introduction to Machine Learning', level=1)
    doc.add_paragraph(
        'Machine learning is a subset of artificial intelligence '
        'that focuses on enabling systems to learn from data '
        'and improve their performance over time without being '
        'explicitly programmed.'
    )
    
    doc.add_heading('Key Concepts', level=2)
    doc.add_paragraph('Supervised Learning', style='List Bullet')
    doc.add_paragraph('Unsupervised Learning', style='List Bullet')
    doc.add_paragraph('Reinforcement Learning', style='List Bullet')
    
    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_content = docx_buffer.getvalue()
    
    filename = "test_ml_doc.docx"
    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    print("=" * 70)
    print("Testing DOCX Upload to RAG System")
    print("=" * 70)
    print()
    
except ImportError:
    print("python-docx not installed, skipping DOCX test")
    print("Install with: uv add --dev python-docx")
    exit(0)

# Upload the DOCX
print(f"Uploading {filename}...")
try:
    files = {"file": (filename, docx_content, content_type)}
    data = {"async_processing": "false"}  # Wait for processing to complete
    
    response = requests.post(
        "http://localhost:8000/rag/upload",
        headers=headers,
        files=files,
        data=data,
        timeout=60
    )
    
    print(f"Status Code: {response.status_code}")
    print()
    
    if response.status_code == 200:
        result = response.json()
        print("✅ DOCX upload successful!")
        print(f"   Document ID: {result.get('document_id')}")
        print(f"   Status: {result.get('status')}")
        print(f"   Chunks Created: {result.get('chunks_created')}")
        print(f"   Processing Time: {result.get('processing_time', 0):.3f}s")
        print(f"   Message: {result.get('message')}")
    else:
        print(f"❌ Upload failed!")
        print(f"   Response: {response.text}")
        
except Exception as e:
    print(f"❌ Error: {e}")
    import traceback
    traceback.print_exc()

print()
print("=" * 70)
